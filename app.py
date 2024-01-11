import streamlit as st
from streamlit_option_menu import option_menu
from docx import Document
from PIL import Image
import datetime
import dateTime
from streamlit.components.v1 import html
def downloadNow_faculty(data):
    file = Document()
    file.add_paragraph("{},".format(data["receiver"]))
    file.add_paragraph("Date:{}".format(data["starting_date"]))
    file.add_paragraph("Subject:{}".format(data["subject"]))
    file.add_paragraph("Sir/Ma'am")
    starting = data["starting_date"]
    end = data["end_date"]
    reason = data["reason"]
    no_of_day = dateTime.dateSubtract(end)
    person_name = data["person_name"]
    Name = data["Name"]
    String = "Sir/Ma'am%0A%0A%0AI am writing to ask you for a {} leave from {} to {} due to {} . I am going to resume work from {}.%0AI shall be reachable on my mobile number and email during the period. My person in charge, {} will be handling my tasks in my absence.%0AI will be thankful to you for considering my application.%0A%0A%0AYours Sincerely,%0A%0A{}".format(no_of_day,starting,end,reason,end,person_name,Name)
    file.add_paragraph("I am writing to ask you for a {} leave from {} to {} due to {} . I am going to resume work from {}.".format(no_of_day,starting,end,reason,end))
    file.add_paragraph("I shall be reachable on my mobile number and email during the period. My person in charge, {} will be handling my tasks in my absence.".format(person_name))
    file.add_paragraph("""I will be thankful to you for considering my application.
    

Yours Sincerely,
    
{}""".format(Name))

    file.save("generatedLetter.docx")
    with open("generatedLetter.docx", "rb") as file:
        btn = st.download_button(
            label="Download Now",
            data=file,
            file_name="generatedLetter.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
          )
    st.markdown('<a href="mailto:{}?subject={}!&body={}" style="background-color: blue;color: white;padding: 8px 12px;text-decoration: none;border-radius: 5px;display: inline-block;">Mail Now!!</a>'.format(data["mail"],data["subject"],String), unsafe_allow_html=True)

def home_faculty():
    name= st.text_input("Name")
    receiver = st.text_area("Receiver Details")
    subject = st.text_input("Enter your subject")
    date1 = str(st.date_input("Enter the starting date"))
    date2 = str(st.date_input("Enter the last date date"))
    reason = st.text_area("Enter The Reason ")
    mail = st.text_input("Enter Email Address")
    person_name = st.text_input("Enter the name of the person who can represent you, during leave")
    if st.button("Submit"):
        datadict = {
        "Name" : name,
        "starting_date" : date1,
        "end_date" : date2,
        "receiver" : receiver,
        "subject" : subject,
        "reason" : reason,
        "person_name" : person_name,
        "mail" : mail
    }
        downloadNow_faculty(datadict)
def letter():
    st.title("Formal Application Generator")
    st.subheader("Dashboard")
    selected = option_menu(
        menu_title="Main Menu",
        options=["Home","Leave Application","About","Contact Us"],
        icons=["house","book","gear","contact"],
        menu_icon="cast",
        orientation="horizontal"
        )
    if selected=="Home":
        st.markdown('''### Hello Everyone,

The formal letter generator is a highly useful tool that can assist individuals in composing a professional and well-crafted formal letter. Whether you need to send a letter to your employer, a university or any other institution, this tool can help you create a letter that is both appropriate and effective.

In today's fast-paced world, time is a valuable commodity and the formal letter generator provides a quick and convenient solution for those who need to send a formal letter. With just a few clicks, you can easily construct a letter that follows the correct tone and style for the intended recipient.

Additionally, the formal letter generator can also help to ensure that all the necessary information is included in the letter, such as the date, recipient's address, and the sender's signature.

In conclusion, the formal letter generator is an invaluable tool for anyone who needs to send a professional and well-written formal letter.

Thank you.

''')
    if selected=="Leave Application":
        home_faculty()
    if selected == "Contact Us":
        st.markdown("""
        ## YOU CAN CONTACT ME -
        Instagram : mrtechnostart(Ram Badan Pandey)                  
        Email Id : rambpandey238@gmail.com
        """)
        st.text("IMS Engineering College, IMS Hostel, Ghaziabad Varanasi")
        st.markdown("""
        
        ## Team Members :-
        1. Ram Badan Pandey (2101430100141)
        2. Rohit Kumar Pandey (2101430100145)
        3. Saif Mohammad (2101430100150)
        4. Subodh Dubey (2101430100177)
        5. Suryansh Katiyar (2101430100180)
        """)
    if selected == "About":
        st.markdown('''### Here is a revised overview of our Python project:

A Formal Application Generator that utilizes the Streamlit library for its frontend interface, and the python-docx library for document generation. This project aims to automate the creation of formal applications and make the process more efficient and streamlined. With Streamlit's user-friendly interface, the application will provide an easy-to-use interface for users to input the required information and generate a professional-looking formal application. The final document will be generated using the python-docx library.

### Goals and Objectives:

To automate the creation of formal applications
To simplify the application creation process
To provide a user-friendly interface for inputting data
To generate professional-looking documents with the help of python-docx


### Features:

A Streamlit-based interface for data input
The ability to generate a professional-looking formal application using python-docx
This project has the potential to greatly streamline the formal application process, saving users time and effort, and resulting in a polished, professional-looking final product.''')
def hideFooter():
    with Image.open("favicon.ico") as icon:
        st.set_page_config(page_title="ApploGen(beta)",page_icon=icon)
    hide_streamlit_style = """
            <style>
            #MainMenu {visibility: hidden;}
            footer {
	            visibility: hidden;
            }
            footer:after {
                content:'Made With ❤️ By MrTechnoStart'; 
                visibility: visible;
                display: block;
                position: relative;
                #background-color: red;
                padding: 5px;
                top: 2px;
            }
                        </style>
                        """
    st.markdown(hide_streamlit_style, unsafe_allow_html=True)
hideFooter()
letter()